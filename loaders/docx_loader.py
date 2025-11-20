"""
DOCX document loader with text extraction
"""
import logging
from pathlib import Path
from typing import List, Dict, Any
from docx import Document

logger = logging.getLogger(__name__)


class DOCXLoader:
    """Load and extract text from DOCX files"""
    
    def __init__(self):
        self.supported_formats = [".docx"]
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Load DOCX file and extract text
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted content and metadata
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_texts.append(row_text)
            
            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\nTables:\n" + "\n".join(table_texts)
            
            metadata = {
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }
            
            # Try to extract document properties
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.created:
                    metadata["created"] = str(core_props.created)
            except:
                pass
            
            return {
                "content": full_text,
                "metadata": metadata,
                "chunks": self._chunk_text(full_text)
            }
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
    
    def _chunk_text(self, text: str, chunk_size: int = 512, overlap: int = 50) -> List[Dict[str, Any]]:
        """
        Split text into chunks for embedding
        
        Args:
            text: Full text content
            chunk_size: Maximum characters per chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks with metadata
        """
        if not text.strip():
            return []
        
        chunks = []
        words = text.split()
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            if current_length + word_length > chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                chunks.append({
                    "text": chunk_text,
                    "start_index": len(" ".join([c["text"] for c in chunks])) if chunks else 0,
                    "length": len(chunk_text)
                })
                
                # Overlap: keep last N words
                overlap_words = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_words
                current_length = sum(len(w) + 1 for w in overlap_words)
            
            current_chunk.append(word)
            current_length += word_length
        
        # Add remaining chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "text": chunk_text,
                "start_index": len(" ".join([c["text"] for c in chunks])),
                "length": len(chunk_text)
            })
        
        return chunks

